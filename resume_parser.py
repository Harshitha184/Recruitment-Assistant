from pypdf import PdfReader
from docx import Document
import re
from langfuse import observe
def clean_text(text):
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return clean_text(text)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""

    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return clean_text(text)

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()
@observe()
def extract_resume_text(file_path):
    ext = file_path.lower()
    if ext.endswith(".pdf"):  
        return extract_text_from_pdf(file_path)
    if ext.endswith(".docx"):  
        return extract_text_from_docx(file_path)
    if ext.endswith(".txt"):   
        return extract_text_from_txt(file_path)
    return ""